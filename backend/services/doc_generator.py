# =============================================================================
# services/doc_generator.py — Generate Word Documents from a Template
# =============================================================================
# WHAT THIS FILE DOES:
#   Takes a Word (.docx) template file and fills in placeholders with
#   actual financial data from the reconciliation results.
#   Produces one .docx file per sector/account.
#
# WHAT IS python-docx?
#   python-docx lets you read and write Microsoft Word files in Python.
#   A .docx file is actually a ZIP file containing XML inside.
#   python-docx handles all the XML complexity for you.
#
# HOW TEMPLATE REPLACEMENT WORKS:
#   In your Word template, you write placeholders like:
#     {{PORTFOLIO_NAME}}   → will be replaced with "Harris Corporation - Large-cap Growth"
#     {{DATE_RANGE}}       → will be replaced with "09/30/2025 - 12/31/2025"
#     {{COMPANY_NAME}}     → will be replaced with each company name
#     {{AVG_WEIGHT}}       → will be replaced with the avg weight value
#     {{PARENT_VALUE}}     → will be replaced with the parent/corrected value
#
#   The template can have a TABLE where each row is filled with one company.
#
# SUPPORTED PLACEHOLDER FORMAT: {{PLACEHOLDER_NAME}}
#   Double curly braces are easy to see in Word and won't conflict with
#   normal text you might write.
#
# IMPORTANT WORD GOTCHA:
#   Word sometimes SPLITS a placeholder across multiple "runs".
#   A "run" is a contiguous piece of text with the same formatting.
#   "{{COMPANY" might be one run, "_NAME}}" might be another run!
#   We handle this by reconstructing the full paragraph text and then
#   replacing it piece by piece.
# =============================================================================

import os
import re
from pathlib import Path
from datetime import datetime
from typing import Optional
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Where to save generated files
OUTPUT_DIR = Path(__file__).parent.parent / "outputs"
OUTPUT_DIR.mkdir(exist_ok=True)


# =============================================================================
# MAIN FUNCTION: Generate Word Files
# =============================================================================
def generate_word_files(
    template_path: str,
    reconciliation_results: dict,
    metadata: dict,
    job_id: int
) -> list:
    """
    Creates one Word file per sector from the reconciliation results.
    
    Args:
        template_path:           Path to the uploaded .docx template
        reconciliation_results:  Output from comparator.reconcile()
        metadata:                Report metadata (portfolio name, dates, etc.)
        job_id:                  The job ID (used in output filename)
    
    Returns:
        List of dicts: [{"filename": "Healthcare.docx", "account_name": "Health care"}, ...]
    """
    # Group results by sector
    sectors = _group_by_sector(reconciliation_results["results"])
    
    generated_files = []
    
    # Generate one Word file per sector
    for sector_name, companies in sectors.items():
        # Create a safe filename (remove special characters)
        safe_sector = re.sub(r'[^\w\s-]', '', sector_name).strip().replace(' ', '_')
        output_filename = f"job{job_id}_{safe_sector}.docx"
        output_path = OUTPUT_DIR / output_filename
        
        # Fill the template and save
        _fill_template(
            template_path=template_path,
            output_path=str(output_path),
            sector_name=sector_name,
            companies=companies,
            metadata=metadata
        )
        
        generated_files.append({
            "filename": output_filename,
            "account_name": sector_name,
            "path": str(output_path)
        })
    
    return generated_files


# =============================================================================
# FUNCTION: Fill Template With Data
# =============================================================================
def _fill_template(
    template_path: str,
    output_path: str,
    sector_name: str,
    companies: list,
    metadata: dict
):
    """
    Opens the template, replaces placeholders, and saves to output_path.
    
    WORD DOCUMENT STRUCTURE (python-docx model):
    Document
    ├── paragraphs[]     ← All text blocks NOT in a table
    └── tables[]         ← All tables in the document
        └── rows[]
            └── cells[]
                └── paragraphs[]
                    └── runs[]   ← Individual text pieces (with formatting)
    
    To replace text while preserving formatting, we work at the paragraph level.
    """
    doc = Document(template_path)
    
    # Build the replacement dictionary
    replacements = {
        "{{PORTFOLIO_NAME}}":  metadata.get("portfolio_name", "Portfolio"),
        "{{BENCHMARK_NAME}}":  metadata.get("benchmark_name", "Benchmark"),
        "{{DATE_RANGE}}":      metadata.get("date_range", ""),
        "{{CURRENCY}}":        metadata.get("currency", "U.S. Dollar"),
        "{{SECTOR_NAME}}":     sector_name,
        "{{GENERATED_DATE}}":  datetime.now().strftime("%B %d, %Y"),
        "{{COMPANY_COUNT}}":   str(len(companies)),
    }
    
    # -------------------------------------------------------------------
    # STEP 1: Replace simple placeholders in all paragraphs
    # -------------------------------------------------------------------
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)
    
    # -------------------------------------------------------------------
    # STEP 2: Replace placeholders in table cells
    # -------------------------------------------------------------------
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)
    
    # -------------------------------------------------------------------
    # STEP 3: Fill company data table (if {{ROW_START}} marker found)
    # -------------------------------------------------------------------
    _fill_company_table(doc, companies)
    
    # Save the completed document
    doc.save(output_path)


# =============================================================================
# HELPER: Replace Text in a Single Paragraph
# =============================================================================
def _replace_in_paragraph(paragraph, replacements: dict):
    """
    Replace placeholder text in a paragraph while preserving formatting.
    
    THE TRICKY PART:
    Word stores paragraph text as multiple "runs" (text segments).
    Each run can have different formatting (bold, italic, font size, color).
    
    When you type {{COMPANY_NAME}} in Word, it MIGHT be stored as:
    Run 1: "{{COMPANY"    (bold, red)
    Run 2: "_NAME}}"      (normal)
    
    If we just replaced Run 1's text, we'd miss the second part.
    
    OUR SOLUTION:
    1. Get the FULL paragraph text (all runs joined together)
    2. Check if any placeholder exists in the full text
    3. If yes, put the replacement text in the FIRST run
       and clear all other runs
    
    This means the replacement uses the FIRST run's formatting.
    """
    # Get the full text of this paragraph
    full_text = paragraph.text
    
    # Check if any placeholder exists
    found_any = False
    for placeholder, replacement in replacements.items():
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, replacement)
            found_any = True
    
    if not found_any:
        return  # Nothing to replace, skip this paragraph
    
    # Replace the text: put everything in run[0], clear the rest
    if paragraph.runs:
        paragraph.runs[0].text = full_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        # No runs exist, add one
        paragraph.add_run(full_text)


# =============================================================================
# HELPER: Fill Company Data Table
# =============================================================================
def _fill_company_table(doc: Document, companies: list):
    """
    Looks for a table that has {{COMPANY_ROW}} in it (a template row marker).
    Duplicates that row for each company and fills in the data.
    
    TEMPLATE TABLE DESIGN:
    | Company Name    | Exchange | Avg Weight | Ending Weight | Return % | Contrib |
    |{{COMPANY_NAME}} |{{EXCHANGE}}|{{AVG_WEIGHT}}|{{END_WEIGHT}}|{{RETURN}}|{{CONTRIB}}|
    
    The second row (with placeholders) is the "template row".
    We duplicate it for each company, fill in the values, then delete the template.
    
    TABLE STRUCTURE IN python-docx:
    - doc.tables  → list of all tables
    - table.rows  → list of all rows
    - row.cells   → list of all cells in that row
    - cell.text   → plain text content of cell
    - cell.paragraphs → access to formatted text
    """
    for table in doc.tables:
        # Find the template row (contains {{COMPANY_NAME}})
        template_row_idx = None
        for i, row in enumerate(table.rows):
            row_text = ' '.join(cell.text for cell in row.cells)
            if "{{COMPANY_NAME}}" in row_text:
                template_row_idx = i
                break
        
        if template_row_idx is None:
            continue  # This table doesn't have a company template row
        
        template_row = table.rows[template_row_idx]
        
        # Add one new row per company (by copying the template row's XML)
        for company in companies:
            # Build field values for this company
            values = {
                "{{COMPANY_NAME}}":    company.get("company_name", ""),
                "{{EXCHANGE}}":        company.get("exchange", ""),
                "{{AVG_WEIGHT}}":      _fmt(company.get("avg_weight")),
                "{{END_WEIGHT}}":      _fmt(company.get("ending_weight")),
                "{{RETURN}}":          _fmt(company.get("return_pct")),
                "{{CONTRIB}}":         _fmt(company.get("contrib")),
                "{{MATCH_SCORE}}":     _fmt(company.get("match_score")),
                "{{STATUS}}":          company.get("status", ""),
                "{{PARENT_AVG_W}}":    _get_parent_val(company, "avg_weight"),
                "{{PARENT_END_W}}":    _get_parent_val(company, "ending_weight"),
                "{{PARENT_RETURN}}":   _get_parent_val(company, "return_pct"),
                "{{PARENT_CONTRIB}}":  _get_parent_val(company, "contrib"),
                "{{DIFF_AVG_W}}":      _get_diff_val(company, "avg_weight"),
                "{{DIFF_RETURN}}":     _get_diff_val(company, "return_pct"),
            }
            
            # Copy the template row by duplicating its XML structure
            # lxml is the XML library underneath python-docx
            import copy
            from lxml import etree
            new_row_xml = copy.deepcopy(template_row._tr)  # _tr = the raw XML element
            
            # Replace placeholders in the copied XML
            xml_str = etree.tostring(new_row_xml, encoding='unicode')
            for placeholder, value in values.items():
                xml_str = xml_str.replace(placeholder, str(value))
            
            # Parse the modified XML back into an lxml element
            new_row_element = etree.fromstring(xml_str)
            
            # Insert the new row after the template row
            template_row._tr.addnext(new_row_element)
        
        # Remove the template row (it served its purpose)
        table._tbl.remove(template_row._tr)


# =============================================================================
# HELPER: Format a numeric value for display
# =============================================================================
def _fmt(val) -> str:
    """Format a value for display in the Word document."""
    if val is None:
        return "N/A"
    if isinstance(val, float):
        return f"{val:.2f}"
    return str(val)


def _get_parent_val(company: dict, field: str) -> str:
    """Get the parent value for a field from the diffs list."""
    for diff in company.get("diffs", []):
        if diff["field"] == field:
            return diff["parent_value"]
    return "N/A"


def _get_diff_val(company: dict, field: str) -> str:
    """Get the diff amount for a field from the diffs list."""
    for diff in company.get("diffs", []):
        if diff["field"] == field:
            val = diff.get("diff_amount")
            return f"{val:+.4f}" if val is not None else "N/A"
    return "N/A"


# =============================================================================
# FUNCTION: Group companies by sector
# =============================================================================
def _group_by_sector(results: list) -> dict:
    """
    Groups a flat list of company results into a dict by sector.
    
    Input:
    [
        {"company_name": "Abbvie", "sector": "Health care", ...},
        {"company_name": "Natera", "sector": "Health care", ...},
        {"company_name": "Apple",  "sector": "Technology", ...},
    ]
    
    Output:
    {
        "Health care": [abbvie_dict, natera_dict],
        "Technology": [apple_dict]
    }
    """
    sectors = {}
    for company in results:
        sector = company.get("sector") or "Unknown Sector"
        if sector not in sectors:
            sectors[sector] = []
        sectors[sector].append(company)
    return sectors


# =============================================================================
# FUNCTION: Create a sample Word template (for testing/demo purposes)
# =============================================================================
def create_sample_template(output_path: str):
    """
    Creates a sample Word template that users can download and customize.
    This shows users exactly what placeholder syntax to use.
    """
    doc = Document()
    
    # Title
    title = doc.add_heading("Performance Attribution Report", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Report info section
    doc.add_paragraph("Portfolio: {{PORTFOLIO_NAME}}")
    doc.add_paragraph("Benchmark: {{BENCHMARK_NAME}}")
    doc.add_paragraph("Date Range: {{DATE_RANGE}}")
    doc.add_paragraph("Currency: {{CURRENCY}}")
    doc.add_paragraph("Generated: {{GENERATED_DATE}}")
    
    doc.add_paragraph()  # Blank line
    
    # Sector heading
    doc.add_heading("Sector: {{SECTOR_NAME}}", level=2)
    doc.add_paragraph(f"Total Companies: {{{{COMPANY_COUNT}}}}")
    
    doc.add_paragraph()
    
    # Company data table
    # Headers
    table = doc.add_table(rows=2, cols=8)
    table.style = "Table Grid"
    
    headers = ["Company", "Ticker", "Child Avg W", "Parent Avg W", 
               "Child Return%", "Parent Return%", "Diff Return%", "Status"]
    placeholders = ["{{COMPANY_NAME}}", "{{EXCHANGE}}", "{{AVG_WEIGHT}}", 
                    "{{PARENT_AVG_W}}", "{{RETURN}}", "{{PARENT_RETURN}}", 
                    "{{DIFF_RETURN}}", "{{STATUS}}"]
    
    # Fill header row
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        # Make header bold
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    
    # Fill template row with placeholders
    template_row = table.rows[1]
    for i, placeholder in enumerate(placeholders):
        template_row.cells[i].text = placeholder
    
    doc.add_paragraph()
    doc.add_paragraph("--- End of Report ---")
    
    doc.save(output_path)
    return output_path
